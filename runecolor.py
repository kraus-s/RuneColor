import pandas as pd
from docx import Document
from docx.shared import RGBColor


# Constants
# ---------

runeCode = "runeDict.xlsx"
inRunes = 'inrunes.txt'
safe = 'a'
unsafe = 'b'

# Colors

unsafe1 = 201, 33, 30   # Rune is safe, variant is not
unsafe2 = 0, 0, 255     # Rune is unsafe, but if given rune is correct reading variant is safe
unsafe3 = 102, 204, 0   # Both are unsafe

# Functions
# ---------

def build_runeDict():
    runePD = pd.read_excel(runeCode, header=None)
    runePD = runePD.rename(columns={0: 'UniRune', 1: 'OutRune'})
    runeDict = runePD.set_index('UniRune')['OutRune'].to_dict()
    return runeDict

runeDict = build_runeDict()

with open(inRunes, 'r', encoding='UTF-8') as iR:
    runeString = iR.read()
    n = 5
    runeList = [runeString[i:i+n] for i in range(0, len(runeString), n)]


outDoc = Document()

# outDoc.add_paragraph('Safe reading of rune, unsafe reading of variant').font.color.rgb=RGBColor(201, 33, 30)
# outDoc.add_paragraph('Unsafe reading of rune, safe guess of variant').font.color.rgb=RGBColor(0, 0, 255)
# outDoc.add_paragraph('Both rune and variant are unsafe').font.color.rgb=RGBColor(102, 204, 0)

p = outDoc.add_paragraph()
for rune in runeList:
    runeValues = [i for i in rune]
    runeType = runeValues[0]+runeValues[2]+runeValues[3]
    runeSafety = runeValues[1]
    typeSafety = runeValues[4]
    gimmeRune = runeDict[runeType]
    if runeSafety == safe:
        if typeSafety == safe:
            print('both safe')
            p.add_run(runeDict[runeType])
        if typeSafety == unsafe:
            print('Rune safe, variant not!')
            p.add_run(runeDict[runeType]).font.color.rgb=RGBColor(201, 33, 30)
        else:
            print("Something went wrong!")
    if runeSafety == unsafe:
        if typeSafety == safe:
            print("Rune unsafe, variant safe")
            p.add_run(runeDict[runeType]).font.color.rgb=RGBColor(0, 0, 255)
        if typeSafety == unsafe:
            print("Bot unsafe")
            p.add_run(runeDict[runeType]).font.color.rgb=RGBColor(102, 204, 0)
        else:
            print("Something went wrong!")

outDoc.save('out.docx')